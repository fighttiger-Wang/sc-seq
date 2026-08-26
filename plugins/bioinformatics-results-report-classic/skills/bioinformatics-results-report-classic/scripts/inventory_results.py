#!/usr/bin/env python3
"""Create a local, shallow inventory of a completed bioinformatics result folder."""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import sys
from collections import Counter
from pathlib import Path
from typing import Any

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")

SUPPORTED = {".xlsx", ".csv", ".tsv", ".png", ".jpg", ".jpeg", ".pdf", ".txt", ".md", ".docx"}
IGNORED_DIRS = {"__pycache__", ".git", ".svn"}

HINTS = [
    ("proportion-roe", re.compile(r"proportion|ratio|roe|composition|fraction|占比|比例", re.I)),
    ("differential-expression", re.compile(r"deg|differential|findmarker|marker|logfc|差异", re.I)),
    ("enrichment", re.compile(r"go|kegg|gsea|gsva|enrich|pathway|富集|通路", re.I)),
    ("cell-communication", re.compile(r"cellchat|nichenet|cellphonedb|ligand|receptor|communication|通讯", re.I)),
    ("trajectory", re.compile(r"pseudotime|trajectory|monocle|slingshot|paga|dpt|velocity|拟时序|轨迹", re.I)),
]


def _ensure_input_folder(path: Path) -> Path:
    resolved = path.resolve()
    if not resolved.is_dir():
        raise ValueError(f"Input folder does not exist: {resolved}")
    return resolved


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _decode_text(path: Path, limit: int = 2400) -> str:
    raw = path.read_bytes()[: limit * 4]
    for encoding in ("utf-8-sig", "gb18030", "utf-16"):
        try:
            return raw.decode(encoding)[:limit]
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")[:limit]


def _tabular_preview(path: Path) -> dict[str, Any]:
    delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
    text = _decode_text(path, 8000)
    rows = list(csv.reader(text.splitlines(), delimiter=delimiter))[:4]
    return {"columns": rows[0][:40] if rows else [], "preview_rows": rows[1:4] if len(rows) > 1 else []}


def _xlsx_preview(path: Path) -> dict[str, Any]:
    try:
        import openpyxl
    except ImportError:
        return {"warning": "openpyxl unavailable; workbook contents not inspected"}
    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sheets = []
    try:
        for worksheet in workbook.worksheets[:30]:
            iterator = worksheet.iter_rows(min_row=1, max_row=3, values_only=True)
            rows = [["" if value is None else str(value) for value in row[:40]] for row in iterator]
            sheets.append({
                "name": worksheet.title,
                "max_row": worksheet.max_row,
                "max_column": worksheet.max_column,
                "columns": rows[0] if rows else [],
                "preview_rows": rows[1:3] if len(rows) > 1 else [],
            })
    finally:
        workbook.close()
    return {"sheets": sheets}


def _image_preview(path: Path) -> dict[str, Any]:
    try:
        from PIL import Image
        with Image.open(path) as image:
            return {"width": image.width, "height": image.height, "mode": image.mode, "format": image.format}
    except Exception as exc:
        return {"warning": f"image metadata unavailable: {exc}"}


def _pdf_preview(path: Path) -> dict[str, Any]:
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        text = ""
        if reader.pages:
            text = (reader.pages[0].extract_text() or "")[:1600]
        return {"pages": len(reader.pages), "first_page_text": text}
    except Exception as exc:
        return {"warning": f"PDF metadata unavailable: {exc}"}


def _docx_preview(path: Path) -> dict[str, Any]:
    try:
        from docx import Document
        document = Document(path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())[:2400]
        return {"text_preview": text, "paragraphs": len(document.paragraphs), "tables": len(document.tables)}
    except Exception as exc:
        return {"warning": f"DOCX preview unavailable: {exc}"}


def _analysis_hints(relative: str) -> list[str]:
    return [name for name, pattern in HINTS if pattern.search(relative)]


def inspect_file(path: Path, root: Path, with_hashes: bool, max_preview_bytes: int) -> dict[str, Any]:
    relative = path.relative_to(root).as_posix()
    suffix = path.suffix.lower()
    stat = path.stat()
    item: dict[str, Any] = {
        "path": relative,
        "extension": suffix,
        "bytes": stat.st_size,
        "analysis_hints": _analysis_hints(relative),
        "role": "existing-report" if suffix == ".html" else "result-source",
    }
    if with_hashes:
        item["sha256"] = _sha256(path)
    if stat.st_size > max_preview_bytes:
        item["content"] = {"warning": f"content preview skipped because file exceeds {max_preview_bytes} bytes"}
        return item
    if suffix in {".csv", ".tsv"}:
        item["content"] = _tabular_preview(path)
    elif suffix == ".xlsx":
        item["content"] = _xlsx_preview(path)
    elif suffix in {".png", ".jpg", ".jpeg"}:
        item["content"] = _image_preview(path)
    elif suffix == ".pdf":
        item["content"] = _pdf_preview(path)
    elif suffix in {".txt", ".md"}:
        item["content"] = {"text_preview": _decode_text(path)}
    elif suffix == ".docx":
        item["content"] = _docx_preview(path)
    return item


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("folder", type=Path)
    parser.add_argument("--output", type=Path, help="Optional local JSON output; stdout is used by default.")
    parser.add_argument("--hashes", action="store_true", help="Calculate SHA-256 for supported files.")
    parser.add_argument("--max-depth", type=int, default=6, help="Maximum relative directory depth to scan (default: 6).")
    parser.add_argument("--max-files", type=int, default=4000, help="Maximum supported files to inventory (default: 4000).")
    parser.add_argument("--max-preview-bytes", type=int, default=50 * 1024 * 1024, help="Skip content previews above this size.")
    args = parser.parse_args()
    try:
        root = _ensure_input_folder(args.folder)
        if args.max_depth < 1 or args.max_files < 1 or args.max_preview_bytes < 1:
            raise ValueError("scan limits must be positive integers")
        candidates = [
            path for path in root.rglob("*")
            if not path.is_symlink()
            and path.is_file()
            and path.suffix.lower() in SUPPORTED | {".html"}
            and not any(part in IGNORED_DIRS for part in path.parts)
            and len(path.relative_to(root).parts) <= args.max_depth
        ]
        paths = sorted(candidates, key=lambda value: value.as_posix().lower())
        truncated = len(paths) > args.max_files
        paths = paths[: args.max_files]
        files = []
        for path in paths:
            try:
                files.append(inspect_file(path, root, args.hashes, args.max_preview_bytes))
            except Exception as exc:
                files.append({
                    "path": path.relative_to(root).as_posix(),
                    "extension": path.suffix.lower(),
                    "bytes": path.stat().st_size,
                    "analysis_hints": _analysis_hints(path.relative_to(root).as_posix()),
                    "role": "existing-report" if path.suffix.lower() == ".html" else "result-source",
                    "inspection_error": str(exc),
                })
        hint_counts = Counter(hint for item in files for hint in item["analysis_hints"])
        payload = {
            "root": str(root),
            "supported_file_count": sum(item["role"] == "result-source" for item in files),
            "existing_report_count": sum(item["role"] == "existing-report" for item in files),
            "scan_truncated": truncated,
            "scan_limits": {"max_depth": args.max_depth, "max_files": args.max_files, "max_preview_bytes": args.max_preview_bytes},
            "analysis_hints": dict(sorted(hint_counts.items())),
            "files": files,
            "notice": "Filename hints are provisional. Confirm groups, sample counts, comparison direction, file roles, scan truncation, and inspection errors before interpretation.",
        }
        encoded = json.dumps(payload, ensure_ascii=False, indent=2)
        if args.output:
            output = args.output.resolve()
            output.parent.mkdir(parents=True, exist_ok=True)
            output.write_text(encoded + "\n", encoding="utf-8")
        else:
            sys.stdout.write(encoded + "\n")
        return 0
    except Exception as exc:
        sys.stderr.write(json.dumps({"error": str(exc)}, ensure_ascii=False) + "\n")
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
